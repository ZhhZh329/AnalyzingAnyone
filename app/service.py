from __future__ import annotations

import csv
import hashlib
import json
import os
import shlex
import shutil as std_shutil
import shutil
import subprocess
import sys
import threading
import zipfile
from datetime import datetime
from io import StringIO
from pathlib import Path
from typing import Any

from .errors import GatewayError, generate_prefixed_id, now_local
from .repository import GatewayRepository
from .schemas import (
    DEFAULT_STAGE_KEYS,
    ArtifactRefs,
    AuxiliaryInfo,
    CreateProjectRequest,
    CreateRunRequest,
    ErrorDetail,
    EvidenceAssemblyView,
    FinalReport,
    Run,
    RunDetail,
    RunStatus,
    StageStatus,
    StageStatusEnum,
)


REPO_ROOT = Path(__file__).resolve().parents[1]
SUPPORTED_EXTENSIONS = {"pdf", "docx", "md", "txt", "xlsx", "csv"}
SOURCE_TYPE_MAPPING = {
    "pdf": "document_pdf",
    "docx": "document_docx",
    "md": "markdown_note",
    "txt": "text_note",
    "xlsx": "spreadsheet_table",
    "csv": "spreadsheet_table",
}
IGNORED_DIRS = {"__MACOSX"}
IGNORED_FILENAMES = {".DS_Store"}
MAX_SOURCE_CHARS = int(os.environ.get("GATEWAY_MAX_SOURCE_CHARS", "12000"))


class GatewayService:
    def __init__(self) -> None:
        self.store_root = Path(os.environ.get("GATEWAY_STORE_ROOT", REPO_ROOT / "gateway_store"))
        self.repo = GatewayRepository(self.store_root)
        self.projects_root = self.repo.projects_root
        self.runs_root = self.repo.runs_root

    def get_analysis_tiers(self) -> dict[str, Any]:
        return {
            "default_tier": "medium",
            "tiers": [
                {
                    "key": "low",
                    "label": "低",
                    "description": "低档：快速模式，优先保证链路稳定与产出速度。",
                    "estimated_lens_count": 35,
                    "estimated_duration_minutes": "10-25",
                },
                {
                    "key": "medium",
                    "label": "中",
                    "description": "中档：标准模式，覆盖核心学科与主要 lenses。",
                    "estimated_lens_count": 80,
                    "estimated_duration_minutes": "30-60",
                },
                {
                    "key": "high",
                    "label": "高",
                    "description": "高档：深度模式，运行更多 lenses 以提升覆盖。",
                    "estimated_lens_count": 211,
                    "estimated_duration_minutes": "90+",
                },
            ],
        }

    def create_project(self, payload: CreateProjectRequest, *, request_id: str) -> dict[str, Any]:
        project_id = generate_prefixed_id("proj")
        subject_id = generate_prefixed_id("subj")
        created_at = now_local().isoformat()
        project = {
            "id": project_id,
            "name": payload.name,
            "description": payload.description,
            "subject": {
                "id": subject_id,
                "display_name": payload.subject.display_name,
                "aliases": payload.subject.aliases,
            },
            "status": "active",
            "created_at": created_at,
            "updated_at": created_at,
            "request_id": request_id,
            "latest_package_id": None,
        }
        project_dir = self._project_dir(project_id)
        (project_dir / "packages").mkdir(parents=True, exist_ok=True)
        self._write_json(project_dir / "project.json", project)
        self.repo.append_request_event(
            project_dir,
            {
                "request_id": request_id,
                "event_type": "create_project",
                "created_at": created_at,
                "project_id": project_id,
                "subject_id": subject_id,
            },
        )
        return {
            "project_id": project_id,
            "subject_id": subject_id,
            "status": "created",
        }

    def ingest_package(
        self,
        project_id: str,
        *,
        request_id: str,
        subject_id: str,
        filename: str,
        package_type: str,
        payload: bytes,
        user_notes: str | None,
        default_source_hint: str | None,
    ) -> tuple[str, dict[str, Any]]:
        project = self._load_project(project_id)
        if project["subject"]["id"] != subject_id:
            raise GatewayError(
                "SUBJECT_PROJECT_MISMATCH",
                "subject_id does not belong to the target project",
                status_code=400,
                details={"project_id": project_id, "subject_id": subject_id},
            )
        if package_type.lower() != "zip" or not filename.lower().endswith(".zip"):
            raise GatewayError(
                "PACKAGE_TYPE_UNSUPPORTED",
                "only zip ingestion packages are supported",
                status_code=400,
                details={"package_type": package_type, "filename": filename},
            )

        package_id = generate_prefixed_id("pkg")
        trace_id = f"trace_{package_id}"
        package_dir = self._package_dir(project_id, package_id)
        original_zip = package_dir / "original.zip"
        extracted_root = package_dir / "extracted"
        input_bundle_dir = package_dir / "input_bundle"
        package_dir.mkdir(parents=True, exist_ok=True)

        sha256 = hashlib.sha256(payload).hexdigest()
        byte_size = len(payload)
        original_zip.write_bytes(payload)

        warnings: list[dict[str, Any]] = []
        raw_files: list[dict[str, Any]] = []
        parsed_files: list[dict[str, Any]] = []
        sources: list[dict[str, Any]] = []

        package_record = {
            "id": package_id,
            "project_id": project_id,
            "subject_id": subject_id,
            "filename": filename,
            "package_type": package_type.lower(),
            "byte_size": byte_size,
            "sha256": sha256,
            "status": "processing",
            "created_at": now_local().isoformat(),
            "updated_at": now_local().isoformat(),
            "request_id": request_id,
            "trace_id": trace_id,
            "user_notes": user_notes,
            "default_source_hint": default_source_hint,
            "raw_file_count": 0,
            "supported_file_count": 0,
            "unsupported_file_count": 0,
            "parsed_success_count": 0,
            "parsed_failed_count": 0,
            "empty_text_count": 0,
            "source_count": 0,
            "source_types": [],
            "warnings": [],
        }
        self._write_json(package_dir / "package.json", package_record)

        try:
            self._extract_zip_safe(original_zip, extracted_root)
            raw_files = self._discover_raw_files(extracted_root)
            parsed_files, sources, warnings = self._parse_files_to_sources(
                package_id=package_id,
                extracted_root=extracted_root,
                raw_files=raw_files,
                default_source_hint=default_source_hint,
            )
            if not sources:
                raise GatewayError(
                    "SOURCE_MATERIALIZATION_FAILED",
                    "no usable source could be materialized from the package",
                    status_code=400,
                    details={
                        "package_id": package_id,
                        "raw_file_count": len(raw_files),
                        "warnings": warnings,
                    },
                )
            self._materialize_input_bundle(
                input_bundle_dir=input_bundle_dir,
                subject=project["subject"]["display_name"],
                sources=sources,
            )
            package_record.update(
                {
                    "status": "completed",
                    "raw_file_count": len(raw_files),
                    "supported_file_count": sum(1 for item in raw_files if item["supported"]),
                    "unsupported_file_count": sum(1 for item in raw_files if not item["supported"]),
                    "parsed_success_count": sum(1 for item in parsed_files if item["parse_status"] == "success"),
                    "parsed_failed_count": sum(1 for item in parsed_files if item["parse_status"] == "failed"),
                    "empty_text_count": sum(1 for item in parsed_files if item["parse_status"] == "empty"),
                    "source_count": len(sources),
                    "source_types": sorted({item["type"] for item in sources}),
                    "warnings": warnings,
                }
            )
        except GatewayError as exc:
            warnings = [*warnings, self._warning_from_exception(exc)]
            package_record.update(
                {
                    "status": "failed",
                    "raw_file_count": len(raw_files),
                    "supported_file_count": sum(1 for item in raw_files if item["supported"]),
                    "unsupported_file_count": sum(1 for item in raw_files if not item["supported"]),
                    "parsed_success_count": sum(1 for item in parsed_files if item["parse_status"] == "success"),
                    "parsed_failed_count": sum(1 for item in parsed_files if item["parse_status"] == "failed"),
                    "empty_text_count": sum(1 for item in parsed_files if item["parse_status"] == "empty"),
                    "source_count": len(sources),
                    "source_types": sorted({item["type"] for item in sources}),
                    "warnings": warnings,
                }
            )
            self._persist_package_artifacts(
                package_dir=package_dir,
                package_record=package_record,
                subject=project["subject"]["display_name"],
                raw_files=raw_files,
                parsed_files=parsed_files,
                sources=sources,
                warnings=warnings,
            )
            raise

        self._persist_package_artifacts(
            package_dir=package_dir,
            package_record=package_record,
            subject=project["subject"]["display_name"],
            raw_files=raw_files,
            parsed_files=parsed_files,
            sources=sources,
            warnings=warnings,
        )

        project["latest_package_id"] = package_id
        project["updated_at"] = now_local().isoformat()
        self._write_json(self._project_dir(project_id) / "project.json", project)
        self.repo.append_request_event(
            package_dir,
            {
                "request_id": request_id,
                "trace_id": trace_id,
                "event_type": "ingest_package",
                "created_at": now_local().isoformat(),
                "project_id": project_id,
                "subject_id": subject_id,
                "package_id": package_id,
                "filename": filename,
                "package_type": package_type.lower(),
            },
        )

        return trace_id, {
            "package_id": package_id,
            "subject_id": subject_id,
            "status": package_record["status"],
            "byte_size": byte_size,
            "sha256": sha256,
            "raw_file_count": package_record["raw_file_count"],
            "supported_file_count": package_record["supported_file_count"],
            "unsupported_file_count": package_record["unsupported_file_count"],
            "parsed_success_count": package_record["parsed_success_count"],
            "parsed_failed_count": package_record["parsed_failed_count"],
            "empty_text_count": package_record["empty_text_count"],
            "source_count": package_record["source_count"],
            "source_types": package_record["source_types"],
            "warnings": package_record["warnings"],
        }
    def _canonical_tier(self, tier_value: str | None) -> str:
        value = (tier_value or "medium").strip().lower()
        mapping = {
            "low": "low",
            "medium": "medium",
            "high": "high",
            "lite": "low",
            "standard": "medium",
            "full": "high",
        }
        return mapping.get(value, "medium")

    def create_run(self, project_id: str, payload: CreateRunRequest, *, request_id: str) -> dict[str, Any]:
        project = self._load_project(project_id)
        if project["subject"]["id"] != payload.subject_id:
            raise GatewayError(
                "SUBJECT_PROJECT_MISMATCH",
                "subject_id does not belong to the target project",
                status_code=400,
                details={"project_id": project_id, "subject_id": payload.subject_id},
            )

        self._ensure_no_active_subject_run(payload.subject_id)
        package = self._select_package(project_id, payload.package_id)

        run_id = generate_prefixed_id("run")
        trace_id = f"trace_{run_id}"
        run_dir = self._run_dir(run_id)
        input_dir = run_dir / "input"
        meta_dir = run_dir / "meta"
        meta_dir.mkdir(parents=True, exist_ok=True)
        shutil.copytree(self._package_dir(project_id, package["id"]) / "input_bundle", input_dir)
        run_input_path = meta_dir / "run_input.json"
        run_feedback_path = meta_dir / "run_feedback.json"
        canonical_tier = self._canonical_tier(payload.run_config.analysis_tier.value)
        self._write_json(
            run_input_path,
            {
                "subject_dir": str(input_dir.resolve()),
                "run_id": run_id,
                "trace_id": trace_id,
                "analysis_tier": canonical_tier,
                "run_config": payload.run_config.model_dump(mode="json"),
            },
        )

        run_record = {
            "id": run_id,
            "project_id": project_id,
            "subject_id": payload.subject_id,
            "package_id": package["id"],
            "request_id": request_id,
            "trace_id": trace_id,
            "subject_display_name": project["subject"]["display_name"],
            "status": RunStatus.QUEUED.value,
            "current_stage": None,
            "started_at": now_local().isoformat(),
            "finished_at": None,
            "updated_at": now_local().isoformat(),
            "output_ref": None,
            "feedback_ref": str(run_feedback_path),
            "input_snapshot_ref": str((input_dir / "manifest.json").resolve()),
            "status_ref": str((run_dir / "status.json").resolve()),
            "stdout_ref": str((meta_dir / "stdout.log").resolve()),
            "stderr_ref": str((meta_dir / "stderr.log").resolve()),
            "run_config": payload.run_config.model_dump(),
        }
        self._write_json(meta_dir / "run.json", run_record)
        status_doc = self._initial_status_document(run_record)
        self._mark_stage_success(status_doc, "input_normalize", output_ref=str(input_dir))
        self._write_json(run_dir / "status.json", status_doc)
        request_events_ref = self.repo.append_request_event(
            meta_dir,
            {
                "request_id": request_id,
                "trace_id": trace_id,
                "event_type": "create_run",
                "created_at": now_local().isoformat(),
                "project_id": project_id,
                "subject_id": payload.subject_id,
                "package_id": package["id"],
                "run_id": run_id,
            },
        )
        run_record["request_events_ref"] = str(request_events_ref.resolve())
        self._write_json(meta_dir / "run.json", run_record)

        worker = threading.Thread(
            target=self._run_workflow_process,
            args=(run_id, run_input_path, run_feedback_path),
            daemon=True,
        )
        worker.start()

        return {
            "run_id": run_id,
            "project_id": project_id,
            "subject_id": payload.subject_id,
            "trace_id": trace_id,
            "status": RunStatus.QUEUED.value,
            "run_config": payload.run_config.model_dump(mode="json"),
            "stage_count": len(DEFAULT_STAGE_KEYS),
        }

    def get_run(self, project_id: str, run_id: str) -> RunDetail:
        run_record = self._load_run_record(run_id)
        if run_record["project_id"] != project_id:
            raise GatewayError(
                "RUN_NOT_FOUND",
                "run does not belong to the target project",
                status_code=404,
                details={"project_id": project_id, "run_id": run_id},
            )
        status_doc = self._read_json(self._run_dir(run_id) / "status.json")
        output_dir = self._resolve_run_output_dir(run_id)
        manifest = self._load_run_manifest(output_dir)
        merged_status = self._merge_status_with_manifest(status_doc, manifest)
        run = Run(
            id=run_record["id"],
            project_id=run_record["project_id"],
            subject_id=run_record["subject_id"],
            trace_id=run_record["trace_id"],
            status=self._normalize_run_status(merged_status.get("status", run_record["status"])),
            current_stage=merged_status.get("current_stage"),
            started_at=merged_status.get("started_at", run_record.get("started_at")),
            finished_at=merged_status.get("finished_at", run_record.get("finished_at")),
            output_ref=str(output_dir) if output_dir else run_record["output_ref"],
        )
        stages = [StageStatus(**stage) for stage in merged_status.get("stages", [])]
        error_payload = merged_status.get("error")
        error = ErrorDetail(**error_payload) if error_payload else None
        artifacts = ArtifactRefs(
            output_dir=str(output_dir) if output_dir else run_record.get("output_ref"),
            run_manifest_ref=self._manifest_ref(output_dir),
            assembly_ref=self._artifact_ref(str(output_dir) if output_dir else run_record.get("output_ref"), "assembly.json"),
            critic_output_ref=self._artifact_ref(str(output_dir) if output_dir else run_record.get("output_ref"), "critic_output.json"),
            synthesis_ref=self._artifact_ref(str(output_dir) if output_dir else run_record.get("output_ref"), "synthesis.json"),
            report_ref=self._artifact_ref(str(output_dir) if output_dir else run_record.get("output_ref"), "report.md"),
            feedback_ref=run_record.get("feedback_ref"),
            stdout_ref=run_record.get("stdout_ref"),
            stderr_ref=run_record.get("stderr_ref"),
        )
        auxiliary = AuxiliaryInfo(
            request_id=run_record.get("request_id"),
            created_at=run_record.get("started_at"),
            updated_at=run_record.get("updated_at") or merged_status.get("updated_at"),
            input_snapshot_ref=run_record.get("input_snapshot_ref"),
            status_ref=run_record.get("status_ref"),
            manifest_ref=self._manifest_ref(output_dir),
            request_events_ref=run_record.get("request_events_ref"),
            latest_feedback_ref=run_record.get("feedback_ref"),
            latest_stdout_ref=run_record.get("stdout_ref"),
            latest_stderr_ref=run_record.get("stderr_ref"),
        )
        return RunDetail(run=run, stages=stages, error=error, artifacts=artifacts, auxiliary=auxiliary)

    def get_evidence_assembly(self, project_id: str, run_id: str) -> dict[str, Any]:
        self._assert_run_project(project_id, run_id)
        artifact = self._resolve_run_output_dir(run_id) / "assembly.json"
        if not artifact.exists():
            self._raise_missing_artifact(run_id, "assembly", artifact)
        data = self._read_json(artifact)
        if not isinstance(data, dict):
            raise GatewayError(
                "ASSEMBLY_FORMAT_INVALID",
                "assembly artifact is not a valid JSON object",
                status_code=500,
                details={"run_id": run_id, "artifact": str(artifact)},
            )
        data.setdefault("subject", "")
        data.setdefault("timeline", [])
        data.setdefault("evidence_cards", [])
        return data

    def get_report(self, project_id: str, run_id: str) -> FinalReport:
        self._assert_run_project(project_id, run_id)
        artifact = self._resolve_run_output_dir(run_id) / "report.md"
        if not artifact.exists():
            self._raise_missing_artifact(run_id, "report", artifact)
        run_record = self._load_run_record(run_id)
        content = artifact.read_text(encoding="utf-8")
        generated_at = datetime.fromtimestamp(artifact.stat().st_mtime).astimezone()
        return FinalReport(
            subject=run_record["subject_display_name"],
            format="markdown",
            content=content,
            generated_at=generated_at,
        )

    def _run_workflow_process(self, run_id: str, run_input_path: Path, run_feedback_path: Path) -> None:
        run_dir = self._run_dir(run_id)
        status_path = run_dir / "status.json"
        stdout_path = run_dir / "meta" / "stdout.log"
        stderr_path = run_dir / "meta" / "stderr.log"
        env = os.environ.copy()
        run_record = self._load_run_record(run_id)
        env["ANALYZINGANYONE_STATUS_PATH"] = str(status_path.resolve())
        env["ANALYZINGANYONE_RUN_ID"] = run_id
        env["ANALYZINGANYONE_TRACE_ID"] = run_record["trace_id"]
        env["ANALYZINGANYONE_RESUME"] = os.environ.get("GATEWAY_WORKFLOW_RESUME", "1")
        self._mark_stage_running(status_path, "assemble")
        self._update_run_record(run_id, status=RunStatus.RUNNING.value, current_stage="assemble")

        override = os.environ.get("GATEWAY_WORKFLOW_COMMAND")
        if override:
            command = shlex.split(override)
        else:
            uv = std_shutil.which("uv")
            command = [uv, "run", "python", "main.py"] if uv else [sys.executable, "main.py"]
        command.extend(
            [
                "--input-file",
                str(run_input_path),
                "--feedback-out",
                str(run_feedback_path),
            ]
        )

        try:
            result = subprocess.run(
                command,
                cwd=REPO_ROOT,
                env=env,
                capture_output=True,
                text=True,
            )
            stdout_path.write_text(result.stdout or "", encoding="utf-8")
            stderr_path.write_text(result.stderr or "", encoding="utf-8")
            if result.returncode != 0:
                self._mark_run_failed(
                    status_path,
                    code="WORKFLOW_PROCESS_FAILED",
                    message="workflow process exited with a non-zero status",
                    details={
                        "returncode": result.returncode,
                        "stderr": (result.stderr or "")[-4000:],
                        "feedback_ref": str(run_feedback_path),
                    },
                )
                self._update_run_record(run_id, status=RunStatus.FAILED.value, finished_at=now_local().isoformat())
                return
            feedback = self._load_feedback(run_feedback_path)
            output_dir = self._normalize_output_dir(feedback["output_dir"])
            self._update_run_record(
                run_id,
                output_ref=str(output_dir),
                finished_at=now_local().isoformat(),
                status=RunStatus.COMPLETED.value if feedback.get("status") == "ok" else RunStatus.FAILED.value,
            )
            if feedback.get("status") != "ok":
                self._mark_run_failed(
                    status_path,
                    code="WORKFLOW_FEEDBACK_ERROR",
                    message=feedback.get("message", "workflow reported a non-ok status"),
                    details=feedback,
                )
                self._update_run_record(run_id, status=RunStatus.FAILED.value, current_stage="report")
                return
            self._ensure_final_status(status_path, output_ref=str(output_dir))
            final_status = self._read_json(status_path).get("status", RunStatus.COMPLETED.value)
            self._update_run_record(run_id, status=final_status, current_stage="report")
        except Exception as exc:  # pragma: no cover
            self._mark_run_failed(
                status_path,
                code="WORKFLOW_PROCESS_EXCEPTION",
                message="failed to launch workflow process",
                details={"exception": str(exc)},
            )
            self._update_run_record(run_id, status=RunStatus.FAILED.value, finished_at=now_local().isoformat())

    def _parse_files_to_sources(
        self,
        *,
        package_id: str,
        extracted_root: Path,
        raw_files: list[dict[str, Any]],
        default_source_hint: str | None,
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
        parsed_files: list[dict[str, Any]] = []
        sources: list[dict[str, Any]] = []
        warnings: list[dict[str, Any]] = []

        for raw_file in raw_files:
            base_metadata = {
                "origin_file_id": raw_file["file_id"],
                "relative_path": raw_file["relative_path"],
                "detected_type": raw_file["detected_type"],
                "byte_size": raw_file["byte_size"],
            }
            if not raw_file["supported"]:
                warning = {
                    "code": "FILE_TYPE_UNSUPPORTED",
                    "message": "file type is not supported for automatic parsing",
                    "details": {"relative_path": raw_file["relative_path"], "detected_type": raw_file["detected_type"]},
                }
                warnings.append(warning)
                parsed_files.append(
                    {
                        "file_id": raw_file["file_id"],
                        "relative_path": raw_file["relative_path"],
                        "detected_type": raw_file["detected_type"],
                        "parse_status": "unsupported",
                        "text_content": "",
                        "metadata": base_metadata,
                        "structure": {},
                    }
                )
                continue

            file_path = extracted_root / raw_file["relative_path"]
            try:
                text_content, metadata, structure = self._parse_supported_file(file_path, raw_file["detected_type"])
            except Exception as exc:
                warning = {
                    "code": "FILE_PARSE_FAILED",
                    "message": "failed to parse file content",
                    "details": {"relative_path": raw_file["relative_path"], "detected_type": raw_file["detected_type"], "error": str(exc)},
                }
                warnings.append(warning)
                parsed_files.append(
                    {
                        "file_id": raw_file["file_id"],
                        "relative_path": raw_file["relative_path"],
                        "detected_type": raw_file["detected_type"],
                        "parse_status": "failed",
                        "text_content": "",
                        "metadata": {**base_metadata, "error": str(exc)},
                        "structure": {},
                    }
                )
                continue

            merged_metadata = {**base_metadata, **metadata}
            normalized_text = text_content.strip()
            if not normalized_text:
                warning = {
                    "code": "FILE_EMPTY_AFTER_PARSE",
                    "message": "file was parsed but no usable text content was extracted",
                    "details": {"relative_path": raw_file["relative_path"], "detected_type": raw_file["detected_type"]},
                }
                warnings.append(warning)
                parsed_files.append(
                    {
                        "file_id": raw_file["file_id"],
                        "relative_path": raw_file["relative_path"],
                        "detected_type": raw_file["detected_type"],
                        "parse_status": "empty",
                        "text_content": "",
                        "metadata": merged_metadata,
                        "structure": structure,
                    }
                )
                continue

            parsed_files.append(
                {
                    "file_id": raw_file["file_id"],
                    "relative_path": raw_file["relative_path"],
                    "detected_type": raw_file["detected_type"],
                    "parse_status": "success",
                    "text_content": normalized_text,
                    "metadata": merged_metadata,
                    "structure": structure,
                }
            )
            sources.extend(
                self._materialize_sources_for_file(
                    package_id=package_id,
                    raw_file=raw_file,
                    text_content=normalized_text,
                    metadata=merged_metadata,
                    default_source_hint=default_source_hint,
                )
            )

        return parsed_files, sources, warnings

    def _materialize_sources_for_file(
        self,
        *,
        package_id: str,
        raw_file: dict[str, Any],
        text_content: str,
        metadata: dict[str, Any],
        default_source_hint: str | None,
    ) -> list[dict[str, Any]]:
        chunks = self._chunk_text(text_content, MAX_SOURCE_CHARS)
        chunk_count = len(chunks)
        source_type = SOURCE_TYPE_MAPPING.get(raw_file["detected_type"], default_source_hint or "text_note")
        title = Path(raw_file["relative_path"]).name
        context = f"package={package_id}; path={raw_file['relative_path']}"
        sources: list[dict[str, Any]] = []

        for index, chunk in enumerate(chunks, start=1):
            source_id = f"src_{raw_file['file_id']}_{index:03d}"
            source_metadata = {
                **metadata,
                "chunk_index": index,
                "chunk_count": chunk_count,
            }
            sources.append(
                {
                    "id": source_id,
                    "type": source_type,
                    "date": "unknown",
                    "title": title,
                    "context": context,
                    "metadata": source_metadata,
                    "content": chunk,
                }
            )

        return sources

    def _materialize_input_bundle(
        self,
        *,
        input_bundle_dir: Path,
        subject: str,
        sources: list[dict[str, Any]],
    ) -> None:
        shutil.rmtree(input_bundle_dir, ignore_errors=True)
        sources_dir = input_bundle_dir / "sources"
        sources_dir.mkdir(parents=True, exist_ok=True)

        manifest_sources: list[dict[str, Any]] = []
        for source in sources:
            relative_file = f"sources/{source['id']}.txt"
            (input_bundle_dir / relative_file).write_text(source["content"], encoding="utf-8")
            manifest_entry = {
                key: value
                for key, value in source.items()
                if key != "content"
            }
            manifest_entry["file"] = relative_file
            manifest_sources.append(manifest_entry)

        manifest = {
            "subject": subject,
            "sources": manifest_sources,
        }
        self._write_json(input_bundle_dir / "manifest.json", manifest)

    def _persist_package_artifacts(
        self,
        *,
        package_dir: Path,
        package_record: dict[str, Any],
        subject: str,
        raw_files: list[dict[str, Any]],
        parsed_files: list[dict[str, Any]],
        sources: list[dict[str, Any]],
        warnings: list[dict[str, Any]],
    ) -> None:
        self._write_json(package_dir / "package.json", package_record)
        self._write_json(package_dir / "raw_files.json", {"files": raw_files})
        self._write_json(package_dir / "parsed_files.json", {"files": parsed_files})
        self._write_json(package_dir / "sources.json", {"subject": subject, "sources": sources})
        self._write_json(package_dir / "warnings.json", {"warnings": warnings})

    def _parse_supported_file(self, file_path: Path, detected_type: str) -> tuple[str, dict[str, Any], dict[str, Any]]:
        if detected_type in {"txt", "md"}:
            return self._parse_text_file(file_path)
        if detected_type == "csv":
            return self._parse_csv_file(file_path)
        if detected_type == "xlsx":
            return self._parse_xlsx_file(file_path)
        if detected_type == "docx":
            return self._parse_docx_file(file_path)
        if detected_type == "pdf":
            return self._parse_pdf_file(file_path)
        raise ValueError(f"unsupported file type: {detected_type}")

    def _parse_text_file(self, file_path: Path) -> tuple[str, dict[str, Any], dict[str, Any]]:
        text_content = file_path.read_text(encoding="utf-8-sig", errors="replace")
        lines = text_content.splitlines()
        return text_content, {"line_count": len(lines)}, {"line_count": len(lines)}

    def _parse_csv_file(self, file_path: Path) -> tuple[str, dict[str, Any], dict[str, Any]]:
        raw_text = file_path.read_text(encoding="utf-8-sig", errors="replace")
        rows = list(csv.reader(StringIO(raw_text)))
        preview_rows = rows[:10]
        rendered_lines = []
        for row in rows:
            clean = [str(cell).strip() for cell in row]
            line = " | ".join(cell for cell in clean if cell)
            if line:
                rendered_lines.append(line)
        text_content = "\n".join(rendered_lines)
        structure = {
            "row_count": len(rows),
            "column_count": max((len(row) for row in rows), default=0),
            "preview_rows": preview_rows,
        }
        metadata = {
            "row_count": len(rows),
            "column_count": structure["column_count"],
        }
        return text_content, metadata, structure

    def _parse_xlsx_file(self, file_path: Path) -> tuple[str, dict[str, Any], dict[str, Any]]:
        try:
            from openpyxl import load_workbook
        except ModuleNotFoundError as exc:  # pragma: no cover
            raise RuntimeError("openpyxl is required to parse xlsx files") from exc

        workbook = load_workbook(file_path, read_only=True, data_only=True)
        try:
            sheet_structures = []
            sections = []
            for sheet in workbook.worksheets:
                preview_rows = []
                rendered_lines = []
                row_count = 0
                column_count = 0
                for row in sheet.iter_rows(values_only=True):
                    values = ["" if cell is None else str(cell).strip() for cell in row]
                    if not any(values):
                        continue
                    row_count += 1
                    column_count = max(column_count, len(values))
                    if len(preview_rows) < 10:
                        preview_rows.append(values)
                    line = " | ".join(value for value in values if value)
                    if line:
                        rendered_lines.append(line)
                if rendered_lines:
                    sections.append(f"# Sheet: {sheet.title}\n" + "\n".join(rendered_lines))
                sheet_structures.append(
                    {
                        "name": sheet.title,
                        "row_count": row_count,
                        "column_count": column_count,
                        "preview_rows": preview_rows,
                    }
                )
            text_content = "\n\n".join(sections)
            metadata = {"sheet_count": len(workbook.worksheets)}
            structure = {"sheets": sheet_structures}
            return text_content, metadata, structure
        finally:
            workbook.close()

    def _parse_docx_file(self, file_path: Path) -> tuple[str, dict[str, Any], dict[str, Any]]:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required to parse docx files") from exc

        document = Document(file_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_previews = []
        table_blocks = []
        for table_index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    rows.append(values)
            if not rows:
                continue
            table_previews.append(rows[:5])
            table_lines = [" | ".join(value for value in row if value) for row in rows]
            table_blocks.append(f"# Table {table_index}\n" + "\n".join(line for line in table_lines if line))

        text_content = "\n\n".join([*paragraphs, *table_blocks])
        metadata = {
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
        }
        structure = {
            "table_count": len(document.tables),
            "table_previews": table_previews[:3],
        }
        return text_content, metadata, structure

    def _parse_pdf_file(self, file_path: Path) -> tuple[str, dict[str, Any], dict[str, Any]]:
        try:
            from pypdf import PdfReader
        except ModuleNotFoundError as exc:  # pragma: no cover
            raise RuntimeError("pypdf is required to parse pdf files") from exc

        reader = PdfReader(str(file_path))
        page_blocks = []
        page_summaries = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            normalized = text.strip()
            page_summaries.append({"page_number": page_index, "char_count": len(normalized)})
            if normalized:
                page_blocks.append(normalized)
        text_content = "\n\n".join(page_blocks)
        metadata = {"page_count": len(reader.pages)}
        structure = {"pages": page_summaries}
        return text_content, metadata, structure

    def _discover_raw_files(self, extracted_root: Path) -> list[dict[str, Any]]:
        raw_files: list[dict[str, Any]] = []
        index = 1
        for path in sorted(extracted_root.rglob("*")):
            if not path.is_file():
                continue
            if path.name in IGNORED_FILENAMES or path.name == "manifest.json":
                continue
            relative_path = path.relative_to(extracted_root)
            if any(part in IGNORED_DIRS for part in relative_path.parts):
                continue
            extension = path.suffix.lower().lstrip(".")
            raw_files.append(
                {
                    "file_id": f"file_{index:03d}",
                    "relative_path": relative_path.as_posix(),
                    "filename": path.name,
                    "extension": extension,
                    "detected_type": extension or "unknown",
                    "byte_size": path.stat().st_size,
                    "supported": extension in SUPPORTED_EXTENSIONS,
                }
            )
            index += 1
        return raw_files

    def _initial_status_document(self, run_record: dict[str, Any]) -> dict[str, Any]:
        now = now_local().isoformat()
        stages = [
            {
                "stage_key": stage_key,
                "status": StageStatusEnum.PENDING.value,
                "started_at": None,
                "finished_at": None,
                "duration_ms": None,
                "error_message": None,
                "output_ref": None,
            }
            for stage_key in DEFAULT_STAGE_KEYS
        ]
        return {
            "run_id": run_record["id"],
            "trace_id": run_record["trace_id"],
            "status": RunStatus.QUEUED.value,
            "current_stage": None,
            "started_at": run_record["started_at"],
            "finished_at": None,
            "updated_at": now,
            "stages": stages,
            "error": None,
        }

    def _ensure_final_status(self, status_path: Path, *, output_ref: str | None = None) -> None:
        status_doc = self._read_json(status_path)
        if status_doc["status"] in {RunStatus.COMPLETED.value, RunStatus.PARTIAL_FAILED.value, RunStatus.FAILED.value}:
            return
        has_partial = any(
            stage["status"] == StageStatusEnum.PARTIAL_FAILED.value
            for stage in status_doc.get("stages", [])
        )
        now = now_local().isoformat()
        stage_started_at = None
        for stage in status_doc.get("stages", []):
            if stage["stage_key"] == "input_normalize":
                continue
            if stage["status"] == StageStatusEnum.PENDING.value:
                stage["started_at"] = stage.get("started_at") or now
                stage["status"] = StageStatusEnum.SUCCESS.value
                stage["finished_at"] = now
                stage["duration_ms"] = self._duration_ms(stage["started_at"], now)
            elif stage["status"] == StageStatusEnum.RUNNING.value:
                stage["status"] = StageStatusEnum.SUCCESS.value
                stage["finished_at"] = now
                stage["duration_ms"] = self._duration_ms(stage["started_at"], now) if stage.get("started_at") else None
            if output_ref is not None and stage["stage_key"] in {"assemble", "discipline", "critique", "synthesize", "report"}:
                stage["output_ref"] = output_ref
            if stage["stage_key"] == "assemble":
                stage_started_at = stage.get("started_at")
        status_doc["status"] = RunStatus.PARTIAL_FAILED.value if has_partial else RunStatus.COMPLETED.value
        status_doc["current_stage"] = "report"
        status_doc["finished_at"] = now
        status_doc["updated_at"] = status_doc["finished_at"]
        self._write_json(status_path, status_doc)

    def _mark_run_failed(self, status_path: Path, *, code: str, message: str, details: dict[str, Any]) -> None:
        status_doc = self._read_json(status_path)
        current_stage = status_doc.get("current_stage") or "input_normalize"
        now = now_local().isoformat()
        for stage in status_doc.get("stages", []):
            if stage["stage_key"] == current_stage:
                stage["status"] = StageStatusEnum.FAILED.value
                stage["finished_at"] = now
                stage["error_message"] = message
                if stage.get("started_at"):
                    stage["duration_ms"] = self._duration_ms(stage["started_at"], now)
                break
        status_doc["status"] = RunStatus.FAILED.value
        status_doc["finished_at"] = now
        status_doc["updated_at"] = now
        status_doc["error"] = {
            "code": code,
            "message": message,
            "details": details,
            "retryable": False,
            "stage_key": current_stage,
        }
        self._write_json(status_path, status_doc)

    def _mark_stage_success(self, status_doc: dict[str, Any], stage_key: str, *, output_ref: str | None = None) -> None:
        now = now_local().isoformat()
        for stage in status_doc.get("stages", []):
            if stage["stage_key"] != stage_key:
                continue
            stage["started_at"] = stage.get("started_at") or now
            stage["status"] = StageStatusEnum.SUCCESS.value
            stage["finished_at"] = now
            stage["duration_ms"] = self._duration_ms(stage["started_at"], now)
            if output_ref is not None:
                stage["output_ref"] = output_ref
            break
        status_doc["updated_at"] = now

    def _mark_stage_running(self, status_path: Path, stage_key: str) -> None:
        status_doc = self._read_json(status_path)
        now = now_local().isoformat()
        for stage in status_doc.get("stages", []):
            if stage["stage_key"] != stage_key:
                continue
            stage["status"] = StageStatusEnum.RUNNING.value
            stage["started_at"] = stage.get("started_at") or now
            break
        status_doc["status"] = RunStatus.RUNNING.value
        status_doc["current_stage"] = stage_key
        status_doc["updated_at"] = now
        self._write_json(status_path, status_doc)

    def _raise_missing_artifact(self, run_id: str, artifact_name: str, artifact_path: Path) -> None:
        status_doc = self._read_json(self._run_dir(run_id) / "status.json")
        run_status = status_doc.get("status", RunStatus.QUEUED.value)
        code = "RUN_OUTPUT_NOT_READY" if run_status in {RunStatus.QUEUED.value, RunStatus.RUNNING.value} else "RUN_OUTPUT_MISSING"
        raise GatewayError(
            code,
            f"{artifact_name} is not available yet",
            status_code=409 if code == "RUN_OUTPUT_NOT_READY" else 404,
            details={
                "run_id": run_id,
                "run_status": run_status,
                "artifact": str(artifact_path),
            },
            retryable=code == "RUN_OUTPUT_NOT_READY",
        )

    def _assert_run_project(self, project_id: str, run_id: str) -> None:
        run_record = self._load_run_record(run_id)
        if run_record["project_id"] != project_id:
            raise GatewayError(
                "RUN_NOT_FOUND",
                "run does not belong to the target project",
                status_code=404,
                details={"project_id": project_id, "run_id": run_id},
            )

    def _ensure_no_active_subject_run(self, subject_id: str) -> None:
        for run_file in self.runs_root.glob("*/meta/run.json"):
            run_record = self._read_json(run_file)
            if run_record["subject_id"] != subject_id:
                continue
            status_path = run_file.parent.parent / "status.json"
            status_doc = self._read_json(status_path) if status_path.exists() else {"status": run_record["status"]}
            if status_doc["status"] in {RunStatus.CREATED.value, RunStatus.QUEUED.value, RunStatus.RUNNING.value}:
                raise GatewayError(
                    "SUBJECT_RUN_CONFLICT",
                    "there is already an active run for this subject",
                    status_code=409,
                    details={"run_id": run_record["id"], "subject_id": subject_id},
                    retryable=True,
                )

    def _select_package(self, project_id: str, package_id: str | None) -> dict[str, Any]:
        if package_id:
            package = self._load_package(project_id, package_id)
        else:
            project = self._load_project(project_id)
            latest = project.get("latest_package_id")
            if not latest:
                raise GatewayError(
                    "PACKAGE_NOT_FOUND",
                    "no validated ingestion package is available for this project",
                    status_code=404,
                    details={"project_id": project_id},
                )
            package = self._load_package(project_id, latest)
        if package["status"] != "completed":
            raise GatewayError(
                "PACKAGE_NOT_READY",
                "the selected package is not ready for run creation",
                status_code=409,
                details={"package_id": package["id"], "status": package["status"]},
            )
        return package

    def _load_feedback(self, run_feedback_path: Path) -> dict[str, Any]:
        if not run_feedback_path.exists():
            raise GatewayError(
                "WORKFLOW_FEEDBACK_MISSING",
                "workflow completed without writing feedback.json",
                status_code=500,
                details={"feedback_path": str(run_feedback_path)},
                retryable=True,
            )
        feedback = self._read_json(run_feedback_path)
        output_dir = feedback.get("output_dir")
        if not output_dir:
            raise GatewayError(
                "WORKFLOW_FEEDBACK_INVALID",
                "workflow feedback is missing output_dir",
                status_code=500,
                details=feedback,
            )
        return feedback

    def _normalize_output_dir(self, raw_output_dir: str) -> Path:
        output_dir = Path(raw_output_dir)
        return output_dir if output_dir.is_absolute() else (REPO_ROOT / output_dir)

    def _resolve_run_output_dir(self, run_id: str) -> Path:
        run_record = self._load_run_record(run_id)
        output_ref = run_record.get("output_ref")
        if output_ref:
            return Path(output_ref)
        feedback_ref = run_record.get("feedback_ref")
        if feedback_ref and Path(feedback_ref).exists():
            feedback = self._load_feedback(Path(feedback_ref))
            output_dir = self._normalize_output_dir(feedback["output_dir"])
            self._update_run_record(run_id, output_ref=str(output_dir))
            return output_dir
        return self._run_dir(run_id) / "output"

    def _load_run_manifest(self, output_dir: Path) -> dict[str, Any] | None:
        manifest_path = output_dir / "run_manifest.json"
        if not manifest_path.exists():
            return None
        manifest = self._read_json(manifest_path)
        return manifest if isinstance(manifest, dict) else None

    def _merge_status_with_manifest(self, status_doc: dict[str, Any], manifest: dict[str, Any] | None) -> dict[str, Any]:
        if not manifest:
            return status_doc

        merged = json.loads(json.dumps(status_doc))
        manifest_status = manifest.get("status")
        if manifest_status:
            merged["status"] = self._normalize_run_status(manifest_status).value
        manifest_started = manifest.get("started_at")
        if manifest_started and not merged.get("started_at"):
            merged["started_at"] = manifest_started
        manifest_finished = manifest.get("finished_at")
        if manifest_finished:
            merged["finished_at"] = manifest_finished
        merged["updated_at"] = now_local().isoformat()

        stage_map = {
            stage["stage_key"]: stage
            for stage in merged.get("stages", [])
            if isinstance(stage, dict) and stage.get("stage_key")
        }
        for raw_stage in manifest.get("stages", []) or []:
            normalized_stage = self._normalize_manifest_stage(raw_stage)
            existing = stage_map.get(normalized_stage["stage_key"])
            if existing:
                existing["status"] = normalized_stage["status"]
                existing["error_message"] = normalized_stage["error_message"]
                if normalized_stage["output_ref"]:
                    existing["output_ref"] = normalized_stage["output_ref"]
            else:
                merged.setdefault("stages", []).append(normalized_stage)
                stage_map[normalized_stage["stage_key"]] = normalized_stage

        error_payload = manifest.get("error")
        if error_payload and not merged.get("error"):
            merged["error"] = {
                "code": error_payload.get("code", "WORKFLOW_ERROR"),
                "message": error_payload.get("message", "workflow reported an error"),
                "details": error_payload.get("details", {}),
                "retryable": bool(error_payload.get("retryable", False)),
                "stage_key": error_payload.get("stage_key"),
            }

        current_stage = merged.get("current_stage")
        if not current_stage:
            running_or_failed = [
                stage["stage_key"]
                for stage in merged.get("stages", [])
                if stage.get("status") in {StageStatusEnum.RUNNING.value, StageStatusEnum.FAILED.value}
            ]
            if running_or_failed:
                merged["current_stage"] = running_or_failed[-1]
            elif merged.get("stages"):
                merged["current_stage"] = merged["stages"][-1]["stage_key"]

        return merged

    def _normalize_manifest_stage(self, raw_stage: dict[str, Any]) -> dict[str, Any]:
        status_value = self._normalize_stage_status(raw_stage.get("status")).value
        started_at = raw_stage.get("started_at")
        finished_at = raw_stage.get("finished_at")
        duration_ms = None
        if started_at and finished_at:
            duration_ms = self._duration_ms(started_at, finished_at)
        return {
            "stage_key": raw_stage.get("stage_key", "unknown"),
            "status": status_value,
            "started_at": started_at,
            "finished_at": finished_at,
            "duration_ms": duration_ms,
            "error_message": raw_stage.get("error_message"),
            "output_ref": raw_stage.get("output_ref"),
        }

    def _normalize_stage_status(self, raw_status: str | None) -> StageStatusEnum:
        mapping = {
            "pending": StageStatusEnum.PENDING,
            "queued": StageStatusEnum.PENDING,
            "running": StageStatusEnum.RUNNING,
            "success": StageStatusEnum.SUCCESS,
            "completed": StageStatusEnum.SUCCESS,
            "ok": StageStatusEnum.SUCCESS,
            "failed": StageStatusEnum.FAILED,
            "error": StageStatusEnum.FAILED,
            "partial_failed": StageStatusEnum.PARTIAL_FAILED,
        }
        if raw_status is None:
            return StageStatusEnum.PENDING
        return mapping.get(str(raw_status).lower(), StageStatusEnum.PENDING)

    def _normalize_run_status(self, raw_status: str | None) -> RunStatus:
        mapping = {
            "created": RunStatus.CREATED,
            "queued": RunStatus.QUEUED,
            "running": RunStatus.RUNNING,
            "completed": RunStatus.COMPLETED,
            "success": RunStatus.COMPLETED,
            "ok": RunStatus.COMPLETED,
            "failed": RunStatus.FAILED,
            "error": RunStatus.FAILED,
            "partial_failed": RunStatus.PARTIAL_FAILED,
            "cancelled": RunStatus.CANCELLED,
        }
        if raw_status is None:
            return RunStatus.QUEUED
        return mapping.get(str(raw_status).lower(), RunStatus.QUEUED)

    def _manifest_ref(self, output_dir: Path | None) -> str | None:
        if output_dir is None:
            return None
        manifest_path = output_dir / "run_manifest.json"
        return str(manifest_path.resolve()) if manifest_path.exists() else None

    def _update_run_record(
        self,
        run_id: str,
        *,
        output_ref: str | None = None,
        finished_at: str | None = None,
        status: str | None = None,
        current_stage: str | None = None,
    ) -> None:
        run_record = self._load_run_record(run_id)
        if output_ref is not None:
            run_record["output_ref"] = output_ref
        if finished_at is not None:
            run_record["finished_at"] = finished_at
        if status is not None:
            run_record["status"] = status
        if current_stage is not None:
            run_record["current_stage"] = current_stage
        run_record["updated_at"] = now_local().isoformat()
        self._write_json(self._run_dir(run_id) / "meta" / "run.json", run_record)

    def _extract_zip_safe(self, archive_path: Path, destination: Path) -> None:
        try:
            with zipfile.ZipFile(archive_path) as zf:
                bad_member = zf.testzip()
                if bad_member:
                    raise GatewayError(
                        "PACKAGE_CORRUPTED",
                        "zip archive is corrupted",
                        status_code=400,
                        details={"bad_member": bad_member},
                    )
                for member in zf.infolist():
                    member_path = Path(member.filename)
                    if member_path.is_absolute() or ".." in member_path.parts:
                        raise GatewayError(
                            "PACKAGE_UNPACK_FAILED",
                            "zip archive contains unsafe paths",
                            status_code=400,
                            details={"member": member.filename},
                        )
                zf.extractall(destination)
        except GatewayError:
            raise
        except zipfile.BadZipFile as exc:
            raise GatewayError(
                "PACKAGE_CORRUPTED",
                "failed to read zip archive",
                status_code=400,
                details={"error": str(exc)},
            ) from exc
        except OSError as exc:
            raise GatewayError(
                "PACKAGE_UNPACK_FAILED",
                "failed to extract zip archive",
                status_code=500,
                details={"error": str(exc)},
                retryable=True,
            ) from exc

    def _warning_from_exception(self, exc: GatewayError) -> dict[str, Any]:
        return {
            "code": exc.code,
            "message": exc.message,
            "details": exc.details,
        }

    def _chunk_text(self, text_content: str, max_chars: int) -> list[str]:
        normalized = text_content.strip()
        if not normalized:
            return []
        if len(normalized) <= max_chars:
            return [normalized]

        paragraphs = [paragraph.strip() for paragraph in normalized.split("\n\n") if paragraph.strip()]
        if not paragraphs:
            return [
                normalized[index:index + max_chars].strip()
                for index in range(0, len(normalized), max_chars)
                if normalized[index:index + max_chars].strip()
            ]

        chunks: list[str] = []
        current = ""
        for paragraph in paragraphs:
            if len(paragraph) > max_chars:
                if current:
                    chunks.append(current)
                    current = ""
                for index in range(0, len(paragraph), max_chars):
                    piece = paragraph[index:index + max_chars].strip()
                    if piece:
                        chunks.append(piece)
                continue
            candidate = paragraph if not current else current + "\n\n" + paragraph
            if len(candidate) <= max_chars:
                current = candidate
            else:
                chunks.append(current)
                current = paragraph
        if current:
            chunks.append(current)
        return chunks

    def _load_project(self, project_id: str) -> dict[str, Any]:
        path = self._project_dir(project_id) / "project.json"
        if not path.exists():
            raise GatewayError(
                "PROJECT_NOT_FOUND",
                "project does not exist",
                status_code=404,
                details={"project_id": project_id},
            )
        return self.repo.load_project(project_id)

    def _load_package(self, project_id: str, package_id: str) -> dict[str, Any]:
        path = self._package_dir(project_id, package_id) / "package.json"
        if not path.exists():
            raise GatewayError(
                "PACKAGE_NOT_FOUND",
                "ingestion package does not exist",
                status_code=404,
                details={"project_id": project_id, "package_id": package_id},
            )
        return self.repo.load_package(project_id, package_id)

    def _load_run_record(self, run_id: str) -> dict[str, Any]:
        path = self._run_dir(run_id) / "meta" / "run.json"
        if not path.exists():
            raise GatewayError(
                "RUN_NOT_FOUND",
                "run does not exist",
                status_code=404,
                details={"run_id": run_id},
            )
        return self.repo.load_run(run_id)

    def _project_dir(self, project_id: str) -> Path:
        return self.projects_root / project_id

    def _package_dir(self, project_id: str, package_id: str) -> Path:
        return self._project_dir(project_id) / "packages" / package_id

    def _run_dir(self, run_id: str) -> Path:
        return self.runs_root / run_id

    def _write_json(self, path: Path, payload: Any) -> None:
        self.repo.write_json(path, payload)

    def _read_json(self, path: Path) -> Any:
        return self.repo.read_json(path)

    def _duration_ms(self, started_at: str, finished_at: str) -> int | None:
        try:
            start = self._coerce_datetime(started_at)
            end = self._coerce_datetime(finished_at)
        except ValueError:
            return None
        return max(int((end - start).total_seconds() * 1000), 0)

    def _coerce_datetime(self, value: str):
        return datetime.fromisoformat(value)

    def _artifact_ref(self, output_dir: str | None, filename: str) -> str | None:
        if not output_dir:
            return None
        return str((Path(output_dir) / filename).resolve())


